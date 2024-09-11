import json
import os
from docx import Document
import sys

# Function to count the number of .docx files in the directory
def count_existing_docs(directory):
    # List all files in the directory
    files = os.listdir(directory)
    # Filter files that end with .docx
    docx_files = [file for file in files if file.endswith(".docx")]
    # Return the count of .docx files
    return len(docx_files)

# Function to create invoice doc from JSON data string
def create_invoice_doc(json_data_str, directory):
    # Parse the JSON string
    data = json.loads(json_data_str)

    document = Document()

    # Add invoice details
    invoice = data.get("invoice", {})
    if invoice.get("client_name"):
        document.add_heading(f"Invoice for {invoice['client_name']}", level=1)
    if invoice.get("invoice_number"):
        document.add_paragraph(f"Invoice Number: {invoice['invoice_number']}")
    if invoice.get("invoice_date"):
        document.add_paragraph(f"Invoice Date: {invoice['invoice_date']}")
    if invoice.get("due_date"):
        document.add_paragraph(f"Due Date: {invoice['due_date']}")

    # Add items with numbering
    document.add_heading("Items", level=2)
    items = data.get("items", [])
    
    for idx, item in enumerate(items, start=1):
        document.add_paragraph(f"Item {idx}:")
        if item.get("description"):
            document.add_paragraph(f"  Description: {item['description']}")
        if item.get("quantity"):
            document.add_paragraph(f"  Quantity: {item['quantity']}")
        if item.get("total_price"):
            document.add_paragraph(f"  Total Price: {item['total_price']}")
        document.add_paragraph(" ")

    # Add subtotal details
    document.add_heading("Subtotal", level=2)
    subtotal = data.get("subtotal", {})
    if subtotal.get("tax"):
        document.add_paragraph(f"Tax: {subtotal['tax']}")
    if subtotal.get("discount"):
        document.add_paragraph(f"Discount: {subtotal['discount']}")
    if subtotal.get("total"):
        document.add_paragraph(f"Total: {subtotal['total']}")

    # Add payment instructions if any
    payment_instructions = data.get("payment_instructions", {})
    if any(payment_instructions.values()):
        document.add_heading("Payment Instructions", level=2)
        if payment_instructions.get("due_date"):
            document.add_paragraph(f"Payment Due Date: {payment_instructions['due_date']}")
        if payment_instructions.get("bank_name"):
            document.add_paragraph(f"Bank Name: {payment_instructions['bank_name']}")
        if payment_instructions.get("account_number"):
            document.add_paragraph(f"Account Number: {payment_instructions['account_number']}")
        if payment_instructions.get("payment_method"):
            document.add_paragraph(f"Payment Method: {payment_instructions['payment_method']}")

    # Count the existing .docx files in the directory
    doc_count = count_existing_docs(directory)

    # Generate a filename based on the number of existing documents
    filename = f"invoice_{doc_count + 1}.docx"
    filepath = os.path.join(directory, filename)

    # Save the document with the generated filename
    document.save(filepath)
    print(f"Invoice document created successfully: {filename}")

# Sample JSON data as string
json_data_str = '{ "name":"John", "age":30, "city":"New York"}' #rf'{sys.argv[1]}'

# Call the function to create the document with the JSON data string
YOUR_FOLDER_PATH=rf"..\OCR-RPA\Document_fill\outputs"
create_invoice_doc(json_data_str,YOUR_FOLDER_PATH)
