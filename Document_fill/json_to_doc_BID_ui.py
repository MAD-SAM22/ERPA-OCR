import json
import os
import sys
import docx
from docx import Document

def count_existing_docs(directory):
    files = os.listdir(directory)
    docx_files = [file for file in files if file.endswith(".docx")]
    return len(docx_files)

def create_invoice_doc(json_data_str, directory):
    data = json.loads(json_data_str)

    document = Document()
    
    # Add personal details
    document.add_heading("Dados Pessoais", level=1)
    document.add_paragraph(f"Nome: {data.get('nome')}")
    document.add_paragraph(f"Filiação: {data.get('filiciao')}")
    document.add_paragraph(f"Data de Expedição: {data.get('data de expedicao')}")
    document.add_paragraph(f"Naturalidade: {data.get('Naturalidade')}")
    document.add_paragraph(f"Data de Nascimento: {data.get('Data de nascimento')}")
    document.add_paragraph(f"CPF: {data.get('CPF')}")

    # Count the existing .docx files in the directory
    doc_count = count_existing_docs(directory)

    # Generate a filename based on the number of existing documents
    filename = f"document_{doc_count + 1}.docx"
    filepath = os.path.join(directory, filename)

    # Save the document with the generated filename
    document.save(filepath)
    print(f"Document created succesfully: {filename}")

# Sample JSON data as string
json_data_str = rf'{sys.argv[1]}'

YOUR_FOLDER_PATH = rf"D:\Grad\OCR-RPA\Document_fill\outputs"
create_invoice_doc(json_data_str, YOUR_FOLDER_PATH)
